"""
ドキュメント処理モジュール
様々な形式のドキュメントを処理してテキストに変換する
"""

import json
import logging
import tempfile
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import io

import pandas as pd
# Text splitting
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_text_splitters import TokenTextSplitter

# Language detection for dynamic chunker selection
from ..utils.translation import TranslationService
from langchain.schema import Document

# ファイル処理のためのインポート
try:
    import pypdf
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..config import config

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """ドキュメント処理クラス"""
    
    def __init__(self):
        # 翻訳サービス（言語検出）
        self.translation_service = TranslationService()

        # デフォルトのフォールバック・スプリッター
        self.default_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", "。", ".", "!", "?", " ", ""]
        )
    
        # サポートされるファイル形式
        self.supported_formats = {
            'txt': self._process_txt_file,
            'pdf': self._process_pdf_file,
            'docx': self._process_docx_file,
            'json': self._process_json_file
        }
    
    def process_uploaded_file(self, file_content: bytes, filename: str) -> List[Document]:
        """
        アップロードされたファイルを処理する
        
        Args:
            file_content: ファイルの内容（バイナリデータ）
            filename: ファイル名
            
        Returns:
            List[Document]: 処理されたドキュメントのリスト
        """
        try:
            # ファイル拡張子を取得
            file_extension = Path(filename).suffix.lower().lstrip('.')
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"サポートされていないファイル形式: {file_extension}")
            
            # ファイル形式に応じて処理
            processor = self.supported_formats[file_extension]
            documents = processor(file_content, filename)
            
            if not documents:
                raise ValueError("ファイルからドキュメントを抽出できませんでした")
            
            logger.info(f"ファイル '{filename}' から {len(documents)} 件のドキュメントを処理しました")
            return documents
            
        except Exception as e:
            logger.error(f"ファイル処理中にエラーが発生しました: {e}")
            raise
    
    def _process_txt_file(self, file_content: bytes, filename: str) -> List[Document]:
        """
        TXTファイルを処理する
        
        Args:
            file_content: ファイルの内容
            filename: ファイル名
            
        Returns:
            List[Document]: 処理されたドキュメントのリスト
        """
        try:
            # テキストファイルをデコード
            text = file_content.decode('utf-8')
            
            # ドキュメントを作成
            document = Document(
                page_content=text,
                metadata={
                    "source": filename,
                    "file_type": "txt",
                    "filename": filename
                }
            )
            
            return [document]
            
        except UnicodeDecodeError:
            # UTF-8で失敗した場合は他のエンコーディングを試す
            try:
                text = file_content.decode('cp932')  # Windows日本語
                document = Document(
                    page_content=text,
                    metadata={
                        "source": filename,
                        "file_type": "txt",
                        "filename": filename,
                        "encoding": "cp932"
                    }
                )
                return [document]
            except UnicodeDecodeError:
                raise ValueError(f"テキストファイルのエンコーディングを読み取れませんでした: {filename}")
    
    def _process_pdf_file(self, file_content: bytes, filename: str) -> List[Document]:
        """
        PDFファイルを処理する
        
        Args:
            file_content: ファイルの内容
            filename: ファイル名
            
        Returns:
            List[Document]: 処理されたドキュメントのリスト
        """
        if not PDF_AVAILABLE:
            raise ValueError("PDFサポートが利用できません。pypdfをインストールしてください。")
        
        try:
            # PDFリーダーを作成
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            documents = []
            
            # 各ページを処理
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                
                if text.strip():  # 空でないページのみ処理
                    document = Document(
                        page_content=text,
                        metadata={
                            "source": filename,
                            "file_type": "pdf",
                            "filename": filename,
                            "page": page_num + 1,
                            "total_pages": len(reader.pages)
                        }
                    )
                    documents.append(document)
            
            return documents
            
        except Exception as e:
            raise ValueError(f"PDFファイルの処理中にエラーが発生しました: {e}")
    
    def _process_docx_file(self, file_content: bytes, filename: str) -> List[Document]:
        """
        DOCXファイルを処理する
        
        Args:
            file_content: ファイルの内容
            filename: ファイル名
            
        Returns:
            List[Document]: 処理されたドキュメントのリスト
        """
        if not DOCX_AVAILABLE:
            raise ValueError("DOCXサポートが利用できません。python-docxをインストールしてください。")
        
        try:
            # DOCXドキュメントを作成
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            # 全段落のテキストを抽出
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 全テキストを結合
            full_text = '\n'.join(paragraphs)
            
            if not full_text.strip():
                raise ValueError("DOCXファイルからテキストを抽出できませんでした")
            
            document = Document(
                page_content=full_text,
                metadata={
                    "source": filename,
                    "file_type": "docx",
                    "filename": filename,
                    "paragraphs": len(paragraphs)
                }
            )
            
            return [document]
            
        except Exception as e:
            raise ValueError(f"DOCXファイルの処理中にエラーが発生しました: {e}")
    
    def _process_json_file(self, file_content: bytes, filename: str) -> List[Document]:
        """
        JSONファイルを処理する
        
        Args:
            file_content: ファイルの内容
            filename: ファイル名
            
        Returns:
            List[Document]: 処理されたドキュメントのリスト
        """
        try:
            # JSONファイルをデコード
            json_text = file_content.decode('utf-8')
            json_data = json.loads(json_text)
            
            # JSONデータを処理
            return self.process_json_documents(json_data, filename)
            
        except json.JSONDecodeError as e:
            raise ValueError(f"JSONファイルの解析中にエラーが発生しました: {e}")
        except UnicodeDecodeError:
            raise ValueError(f"JSONファイルのエンコーディングを読み取れませんでした: {filename}")
    
    def get_supported_formats(self) -> List[str]:
        """
        サポートされるファイル形式のリストを返す
        
        Returns:
            List[str]: サポートされるファイル形式のリスト
        """
        formats = ['txt', 'json']
        
        if PDF_AVAILABLE:
            formats.append('pdf')
        
        if DOCX_AVAILABLE:
            formats.append('docx')
        
        return formats
    
    def process_json_documents(self, json_data: Union[List[Dict[str, Any]], Dict[str, Any]], source_filename: str = None) -> List[Document]:
        """
        JSONデータからドキュメントを処理する
        
        Args:
            json_data: JSON形式のドキュメントデータ
            source_filename: ソースファイル名（オプション）
            
        Returns:
            List[Document]: 処理されたドキュメントのリスト
        """
        documents = []
        
        # データが辞書の場合はリストに変換
        if isinstance(json_data, dict):
            json_data = [json_data]
        
        for idx, item in enumerate(json_data):
            try:
                # JSONアイテムからテキストを抽出
                text = self._extract_text_from_json_item(item)
                
                if text:
                    # メタデータを作成
                    metadata = {
                        "source": source_filename or f"document_{idx}",
                        "doc_id": idx,
                        "file_type": "json"
                    }
                    
                    # ファイル名が指定されている場合は追加
                    if source_filename:
                        metadata["filename"] = source_filename
                    
                    # その他のメタデータを追加
                    for k, v in item.items():
                        if k != "content" and isinstance(v, (str, int, float, bool)):
                            metadata[k] = v
                    
                    # ドキュメントを作成
                    doc = Document(page_content=text, metadata=metadata)
                    documents.append(doc)
                    
            except Exception as e:
                logger.error(f"JSON項目の処理中にエラーが発生しました（インデックス: {idx}）: {e}")
                continue
        
        logger.info(f"処理された文書数: {len(documents)}")
        return documents
    
    def _extract_text_from_json_item(self, item: Dict[str, Any]) -> str:
        """
        JSON項目からテキストを抽出する
        
        Args:
            item: JSON項目
            
        Returns:
            str: 抽出されたテキスト
        """
        text_parts = []
        
        # 一般的なテキストフィールドを検索
        text_fields = ["content", "text", "body", "description", "summary", "title", "question", "answer"]
        
        for field in text_fields:
            if field in item and isinstance(item[field], str):
                text_parts.append(item[field])
        
        # その他の文字列値も追加
        for key, value in item.items():
            if key not in text_fields and isinstance(value, str) and len(value) > 10:
                text_parts.append(f"{key}: {value}")
        
        return "\n".join(text_parts)
    
    def _get_splitter_for_language(self, lang_code: str):
        """言語コードに応じたセマンティックチャンク用スプリッターを返す"""
        try:
            # langchain >=0.2 のAPI: RecursiveCharacterTextSplitter.from_language
            splitter = RecursiveCharacterTextSplitter.from_language(
                language=lang_code,
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap
            )
            return splitter
        except Exception:
            # 未サポート言語やエラー時はデフォルトにフォールバック
            return self.default_splitter

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """ドキュメントを言語に応じてセマンティックチャンクに分割する"""
        split_docs: List[Document] = []
        try:
            for doc in documents:
                # 言語検出（短すぎるテキストは検出しない）
                sample_text = doc.page_content[:1000]
                detected_lang = self.translation_service.detect_language(sample_text) or "en"

                splitter = self._get_splitter_for_language(detected_lang)
                chunks = splitter.split_documents([doc])
                split_docs.extend(chunks)

            logger.info(f"分割された文書数: {len(split_docs)} (元:{len(documents)})")
            return split_docs
        except Exception as e:
            logger.error(f"ドキュメント分割中にエラーが発生しました: {e}")
            # エラー時はフォールバックとして元文書を返す
            return documents
    
    def load_json_from_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        ファイルからJSONデータを読み込む
        
        Args:
            file_path: JSONファイルのパス
            
        Returns:
            List[Dict[str, Any]]: JSON データ
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # データが辞書の場合はリストに変換
            if isinstance(data, dict):
                data = [data]
            
            logger.info(f"JSONファイルから{len(data)}件の項目を読み込みました")
            return data
            
        except Exception as e:
            logger.error(f"JSONファイルの読み込み中にエラーが発生しました: {e}")
            return []
    
    def validate_documents(self, documents: List[Document]) -> List[Document]:
        """
        ドキュメントの有効性を検証する
        
        Args:
            documents: 検証対象のドキュメント
            
        Returns:
            List[Document]: 有効なドキュメントのリスト
        """
        valid_docs = []
        
        for doc in documents:
            # 空のドキュメントを除外
            if doc.page_content and len(doc.page_content.strip()) > 0:
                valid_docs.append(doc)
        
        logger.info(f"有効なドキュメント数: {len(valid_docs)}")
        return valid_docs 